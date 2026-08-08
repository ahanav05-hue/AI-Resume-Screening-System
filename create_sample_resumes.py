from docx import Document

def create_resume(filename, name, title, skills, experience, education):

    doc = Document()

    doc.add_heading(name, level=1)
    doc.add_heading(title, level=2)

    doc.add_heading("Professional Summary", level=2)
    doc.add_paragraph(
        "Motivated software professional with strong programming, problem-solving, "
        "and software development skills."
    )

    doc.add_heading("Technical Skills", level=2)

    for skill in skills:
        doc.add_paragraph(skill, style="List Bullet")

    doc.add_heading("Experience", level=2)
    doc.add_paragraph(experience)

    doc.add_heading("Education", level=2)
    doc.add_paragraph(education)

    doc.save(filename)


create_resume(
    "sample_resumes/Software_Engineer_Resume.docx",
    "Rahul Sharma",
    "Software Engineer",
    [
        "Python",
        "Java",
        "SQL",
        "Git",
        "GitHub",
        "Flask",
        "REST API",
        "Linux",
        "Docker"
    ],
    "Developed scalable backend applications using Python and Flask.",
    "B.E. Computer Science"
)

create_resume(
    "sample_resumes/Python_Developer_Resume.docx",
    "Priya Nair",
    "Python Developer",
    [
        "Python",
        "Flask",
        "SQL",
        "HTML",
        "CSS",
        "JavaScript",
        "Git",
        "OOP"
    ],
    "Built web applications using Flask and SQLite.",
    "B.Tech Information Technology"
)

create_resume(
    "sample_resumes/Data_Analyst_Resume.docx",
    "Arjun Kumar",
    "Data Analyst",
    [
        "Python",
        "Pandas",
        "NumPy",
        "SQL",
        "Statistics",
        "Machine Learning",
        "Excel"
    ],
    "Performed business analytics and dashboard reporting.",
    "B.Sc Data Science"
)

print("✅ Sample resumes created successfully!")